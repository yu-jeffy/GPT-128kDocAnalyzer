import os
import csv
from docx import Document
import openai
from dotenv import load_dotenv

# sqlite fix for python 3.10.2
import pysqlite3
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')

from langchain.llms import OpenAI
from langchain.embeddings import OpenAIEmbeddings
from langchain.document_loaders import PyPDFLoader
from langchain.schema import Document as LangchainDocument
from langchain.vectorstores import Chroma
from langchain.agents.agent_toolkits import (
    create_vectorstore_agent,
    VectorStoreToolkit,
    VectorStoreInfo
)

load_dotenv()
api_key = os.getenv('OPENAI_API_KEY')
openai.api_key = api_key

# Function to extract text from various file types
def extract_text(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""
    if file_extension == '.pdf':
        loader = PyPDFLoader(file_path)
        documents = loader.load_and_split()
        # Extract the page_content from each Document object
        text = "\n".join([doc.page_content for doc in documents])
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    elif file_extension == '.csv':
        with open(file_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            text = "\n".join([", ".join(row) for row in reader])
    elif file_extension == '.docx':
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file type. Supported types are: PDF, TXT, CSV, DOCX")
    return text

# Ask the user for the file path and analysis type
file_path = input('Please enter the file path to the document: ')
analysis_type = input('Choose analysis type ("plaintext" or "vector"): ')

# Extract text from the file
document_text = extract_text(file_path)

# Analysis logic based on user choice
if analysis_type == 'plaintext':
    # Direct plaintext analysis
    user_prompt = input('Input your prompt here: ')
    full_prompt = document_text + "\n" + user_prompt
    response = openai.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": full_prompt}
        ]
    )
    # Print just the content of the response
    print(response.choices[0].message.content)
elif analysis_type == 'vector':
    # Vector-based analysis
    embeddings = OpenAIEmbeddings()

    # Instantiate the language model with the API key
    llm = OpenAI(api_key=api_key)

    # Create a Langchain Document object with the extracted text
    langchain_document = LangchainDocument(page_content=document_text, metadata={"source": "User-provided document"})

    # Create a vector store from the Langchain Document object
    store = Chroma.from_documents([langchain_document], embeddings, collection_name='document_collection')
    vectorstore_info = VectorStoreInfo(
        name="document_vectorstore",
        description="Vector store for the document",
        vectorstore=store
    )
    toolkit = VectorStoreToolkit(vectorstore_info=vectorstore_info)
    agent_executor = create_vectorstore_agent(llm=llm, toolkit=toolkit, verbose=True)
    user_prompt = input('Input your prompt here: ')
    
    # Run the agent executor and capture the response
    response = agent_executor.run(user_prompt)
    
    # Assuming the final answer is in the 'Final Answer' section of the response
    # Extract and print the final answer
    final_answer = response.get('Final Answer')
    if final_answer:
        print(final_answer)
    else:
        # If there is no 'Final Answer' key, print the whole response or handle it as needed
        print(response)
else:
    print("Invalid analysis type selected.")